"""
AI Templates Routes
Handles AI template browsing, execution, streaming, and export
"""

from flask import Blueprint, jsonify, request, Response, send_file, session
from app.core.time_utils import utcnow_naive
import json
import time
from jinja2 import Template as Jinja2Template
from io import BytesIO
import markdown

from app.database import db
from app.models.researcher.ai_templates import (
    AITemplate, AIWorkflowExecution, AIWorkbook, WorkbookDocument
)
from app.routes.route_entity_lookup import get_entity_or_404

ai_templates_bp = Blueprint('ai_templates', __name__, url_prefix='/researcher/ai')


def _get_template_or_404(template_id):
    return get_entity_or_404(AITemplate, template_id)


def _get_execution_or_404(execution_id):
    return get_entity_or_404(AIWorkflowExecution, execution_id)


@ai_templates_bp.route('/templates/<template_id>')
def get_template(template_id):
    """Get template configuration"""
    # Try by ID first
    if template_id.isdigit():
        template = _get_template_or_404(int(template_id))
    else:
        # Try by slug/identifier
        template = AITemplate.query.filter_by(identifier=template_id).first_or_404()
    
    return jsonify({
        'id': template.id,
        'name': template.name,
        'description': template.description,
        'category': template.category,
        'input_schema': template.input_schema,
        'prompt_template': template.prompt_template,
        'creativity_default': template.creativity_default
    })


@ai_templates_bp.route('/execute', methods=['POST'])
def execute_template():
    """Create an execution record and return execution ID"""
    data = request.get_json()
    
    template_id = data.get('template_id')
    inputs = data.get('inputs', {})
    
    template = _get_template_or_404(template_id)
    
    # Create execution record
    execution = AIWorkflowExecution(
        template_id=template.id,
        user_id=session.get('user_id'),
        input_data=inputs,
        status='pending'
    )
    
    db.session.add(execution)
    db.session.commit()
    
    return jsonify({
        'execution_id': execution.id,
        'status': 'pending'
    })


@ai_templates_bp.route('/stream/<int:execution_id>')
def stream_generation(execution_id):
    """Server-Sent Events stream for real-time generation"""
    execution = _get_execution_or_404(execution_id)
    
    def generate():
        try:
            # Update status to running
            execution.status = 'running'
            execution.started_at = utcnow_naive()
            db.session.commit()
            
            # Send initial progress
            yield f"event: progress\ndata: {json.dumps({'progress': 0, 'message': 'Preparing prompt...'})}\n\n"
            
            # Render prompt template with inputs
            template_obj = execution.template
            prompt_template = Jinja2Template(template_obj.prompt_template)
            rendered_prompt = prompt_template.render(**execution.input_data)
            
            # Send progress
            yield f"event: progress\ndata: {json.dumps({'progress': 20, 'message': 'Connecting to AI...'})}\n\n"
            
            # Call AI service (integration with Beep.AI.Server)
            from app.services.ai_service import generate_with_llm
            
            full_response = ""
            token_count = 0
            
            # Stream tokens from LLM
            for chunk in generate_with_llm(
                prompt=rendered_prompt,
                creativity=execution.input_data.get('_creativity', 70) / 100.0,
                max_tokens=execution.input_data.get('_maxLength', 500) * 2,  # Rough conversion
                stream=True
            ):
                token = chunk.get('token', '')
                full_response += token
                token_count += 1
                
                # Send token
                yield f"event: token\ndata: {json.dumps({'token': token})}\n\n"
                
                # Send progress periodically
                if token_count % 50 == 0:
                    progress = min(20 + (token_count / 10), 95)
                    yield f"event: progress\ndata: {json.dumps({'progress': progress, 'message': 'Generating...'})}\n\n"
                
                time.sleep(0.01)  # Throttle for smoother streaming
            
            # Save result
            execution.output_data = {'text': full_response}
            execution.result_text = full_response
            execution.tokens_used = token_count
            execution.status = 'completed'
            execution.completed_at = utcnow_naive()
            
            db.session.commit()
            
            # Send completion
            completion_data = json.dumps({
                'success': True,
                'tokens': token_count,
                'model': 'GPT-4'
            })
            yield f"event: complete\ndata: {completion_data}\n\n"
            
        except Exception as e:
            execution.status = 'failed'
            execution.error_message = str(e)
            db.session.commit()
            
            error_data = json.dumps({'error': str(e)})
            yield f"event: error\ndata: {error_data}\n\n"
    
    return Response(generate(), mimetype='text/event-stream')


@ai_templates_bp.route('/export/<int:execution_id>/<format>')
def export_result(execution_id, format):
    """Export execution result in various formats"""
    execution = _get_execution_or_404(execution_id)
    
    if not execution.result_text:
        return jsonify({'error': 'No result to export'}), 400
    
    content = execution.result_text
    template_name = execution.template.name.replace(' ', '_')
    
    if format == 'markdown':
        # Return as markdown file
        buffer = BytesIO(content.encode('utf-8'))
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'{template_name}.md',
            mimetype='text/markdown'
        )
    
    elif format == 'docx':
        # Convert to Word document
        from docx import Document
        from docx.shared import Inches, Pt
        
        doc = Document()
        doc.add_heading(execution.template.name, 0)
        doc.add_paragraph(content)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'{template_name}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    elif format == 'pdf':
        # Convert to PDF
        from weasyprint import HTML
        
        html_content = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; padding: 40px; }}
                h1 {{ color: #1E40AF; }}
            </style>
        </head>
        <body>
            <h1>{execution.template.name}</h1>
            {markdown.markdown(content)}
        </body>
        </html>
        """
        
        buffer = BytesIO()
        HTML(string=html_content).write_pdf(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'{template_name}.pdf',
            mimetype='application/pdf'
        )
    
    else:
        return jsonify({'error': 'Invalid format'}), 400


@ai_templates_bp.route('/workbooks')
def get_workbooks():
    """Get user's workbooks"""
    user_id = session.get('user_id')
    
    workbooks = AIWorkbook.query.filter_by(user_id=user_id).all()
    
    return jsonify([
        {
            'id': wb.id,
            'name': wb.name,
            'description': wb.description,
            'document_count': WorkbookDocument.query.filter_by(workbook_id=wb.id).count(),
            'created_at': wb.created_at.isoformat() if wb.created_at else None
        }
        for wb in workbooks
    ])


@ai_templates_bp.route('/save-to-workbook', methods=['POST'])
def save_to_workbook():
    """Save execution result to workbook"""
    data = request.get_json()
    
    execution_id = data.get('execution_id')
    workbook_id = data.get('workbook_id')
    new_workbook_name = data.get('new_workbook_name')
    document_title = data.get('document_title')
    notes = data.get('notes', '')
    
    execution = _get_execution_or_404(execution_id)
    
    # Create new workbook if needed
    if not workbook_id and new_workbook_name:
        workbook = AIWorkbook(
            name=new_workbook_name,
            user_id=session.get('user_id'),
            project_id=data.get('project_id')
        )
        db.session.add(workbook)
        db.session.flush()
        workbook_id = workbook.id
    
    if not workbook_id:
        return jsonify({'success': False, 'error': 'No workbook specified'}), 400
    
    # Create document
    document = WorkbookDocument(
        workbook_id=workbook_id,
        title=document_title,
        content=execution.result_text,
        source_type='template_execution',
        template_id=execution.template_id,
        execution_id=execution.id,
        notes=notes
    )
    
    db.session.add(document)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'document_id': document.id,
        'workbook_id': workbook_id
    })


@ai_templates_bp.route('/browse')
def browse_templates():
    """Browse all available templates"""
    category = request.args.get('category')
    
    query = AITemplate.query.filter_by(is_system=True)
    
    if category:
        query = query.filter_by(category=category)
    
    templates = query.order_by(AITemplate.usage_count.desc()).all()
    
    return jsonify([
        {
            'id': t.id,
            'name': t.name,
            'description': t.description,
            'category': t.category,
            'icon': t.icon,
            'usage_count': t.usage_count or 0
        }
        for t in templates
    ])
