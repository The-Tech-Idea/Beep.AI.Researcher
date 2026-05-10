"""Admin document management service."""
from __future__ import annotations

import io
import uuid
import hashlib
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from sqlalchemy import func

from app.core.time_utils import utcnow_naive
from app.database import db


@dataclass(frozen=True)
class DocumentListResult:
    """Paginated admin document list result."""

    pagination: Any
    total_count: int
    total_storage: int


class DocumentManagerService:
    """Document operations used by admin routes."""

    allowed_upload_extensions = {
        ".pdf", ".txt", ".text", ".md", ".html", ".htm", ".docx", ".doc", ".pptx",
        ".xlsx", ".xls", ".csv", ".json", ".png", ".jpg", ".jpeg", ".tif", ".tiff",
        ".bmp", ".webp", ".wav", ".mp3", ".m4a", ".flac", ".ogg", ".webm", ".vtt",
        ".srt", ".tex", ".xml", ".xbrl",
    }

    def search_documents(
        self,
        *,
        page: int = 1,
        per_page: int = 50,
        user_id: int | None = None,
        tenant_id: int | None = None,
        project_id: int | None = None,
        search: str = "",
        status: str = "",
        file_type: str = "",
        date_from: str = "",
        date_to: str = "",
        sort: str = "created_desc",
    ) -> DocumentListResult:
        from app.models.researcher import ResearchProject, ResearcherDocument

        query = ResearcherDocument.query

        joined_project = False
        needs_project_join = user_id or tenant_id or sort in {"owner", "project"}
        if user_id:
            query = query.join(ResearchProject).filter(ResearchProject.owner_id == user_id)
            joined_project = True
        if tenant_id:
            if not joined_project:
                query = query.join(ResearchProject)
                joined_project = True
            query = query.filter(ResearchProject.tenant_id == tenant_id)
        if project_id:
            query = query.filter(ResearcherDocument.project_id == project_id)
        if search:
            query = query.filter(ResearcherDocument.filename.ilike(f"%{search}%"))
        if status:
            query = query.filter(ResearcherDocument.status == status)
        if file_type:
            suffix = file_type.strip().lower().lstrip(".")
            query = query.filter(ResearcherDocument.filename.ilike(f"%.{suffix}"))
        if date_from:
            try:
                from datetime import datetime
                query = query.filter(ResearcherDocument.created_at >= datetime.fromisoformat(date_from))
            except ValueError:
                pass
        if date_to:
            try:
                from datetime import datetime, time
                end_date = datetime.combine(datetime.fromisoformat(date_to).date(), time.max)
                query = query.filter(ResearcherDocument.created_at <= end_date)
            except ValueError:
                pass

        if needs_project_join and not joined_project:
            query = query.join(ResearchProject)
            joined_project = True

        if sort == "created_asc":
            order_by = ResearcherDocument.created_at.asc()
        elif sort == "size_desc":
            order_by = ResearcherDocument.file_size.desc()
        elif sort == "size_asc":
            order_by = ResearcherDocument.file_size.asc()
        elif sort == "owner":
            order_by = ResearchProject.owner_id.asc()
        elif sort == "project":
            order_by = ResearchProject.name.asc()
        elif sort == "status":
            order_by = ResearcherDocument.status.asc()
        else:
            order_by = ResearcherDocument.created_at.desc()

        pagination = query.order_by(order_by).paginate(
            page=page,
            per_page=per_page,
            error_out=False,
        )
        total_count = ResearcherDocument.query.count()
        total_storage = (
            db.session.query(func.coalesce(func.sum(ResearcherDocument.file_size), 0)).scalar()
            or 0
        )
        return DocumentListResult(
            pagination=pagination,
            total_count=total_count,
            total_storage=int(total_storage),
        )

    def search_user_documents(
        self,
        *,
        user_id: int,
        page: int = 1,
        per_page: int = 50,
        project_id: int | None = None,
        search: str = "",
        status: str = "",
        file_type: str = "",
    ) -> DocumentListResult:
        """Search documents visible to one user across owned/member projects."""

        from app.models.researcher import ProjectMember, ResearchProject, ResearcherDocument

        visible_project_ids = self._visible_project_ids(user_id)
        query = ResearcherDocument.query.filter(ResearcherDocument.project_id.in_(visible_project_ids))
        if project_id and project_id in visible_project_ids:
            query = query.filter(ResearcherDocument.project_id == project_id)
        elif project_id:
            query = query.filter(False)
        if search:
            query = query.filter(ResearcherDocument.filename.ilike(f"%{search}%"))
        if status:
            query = query.filter(ResearcherDocument.status == status)
        if file_type:
            suffix = file_type.strip().lower().lstrip(".")
            query = query.filter(ResearcherDocument.filename.ilike(f"%.{suffix}"))

        pagination = query.order_by(ResearcherDocument.created_at.desc()).paginate(
            page=page,
            per_page=per_page,
            error_out=False,
        )
        total_count = ResearcherDocument.query.filter(
            ResearcherDocument.project_id.in_(visible_project_ids)
        ).count()
        total_storage = (
            db.session.query(func.coalesce(func.sum(ResearcherDocument.file_size), 0))
            .filter(ResearcherDocument.project_id.in_(visible_project_ids))
            .scalar()
            or 0
        )
        return DocumentListResult(
            pagination=pagination,
            total_count=total_count,
            total_storage=int(total_storage),
        )

    def list_user_projects(self, user_id: int):
        """Return projects where the user is owner or member."""

        from app.models.researcher import ResearchProject

        visible_project_ids = self._visible_project_ids(user_id)
        if not visible_project_ids:
            return []
        return (
            ResearchProject.query
            .filter(ResearchProject.id.in_(visible_project_ids))
            .order_by(ResearchProject.name)
            .all()
        )

    def upload_user_document(self, *, user_id: int, project_id: int, file_storage) -> dict[str, Any]:
        """Store a user upload, create the document record, and sync text to RAG."""

        from app.models.researcher import ResearchProject, ResearcherDocument
        from app.services.quota_service import QuotaExceededError, quota_service
        from app.services.storage import get_storage_backend

        visible_project_ids = self._visible_project_ids(user_id)
        if project_id not in visible_project_ids:
            raise PermissionError("Project is not available to this user.")

        project = db.session.get(ResearchProject, project_id)
        if project is None:
            raise LookupError("Project not found.")
        if not file_storage or not file_storage.filename:
            raise ValueError("Choose a document to upload.")

        filename = Path(file_storage.filename).name
        extension = Path(filename).suffix.lower()
        if extension not in self.allowed_upload_extensions:
            raise ValueError(f"File type {extension or '(none)'} is not allowed.")

        raw_bytes = file_storage.read()
        if not raw_bytes:
            raise ValueError("The uploaded file is empty.")

        try:
            quota_service.check_quota(
                user_id=user_id,
                upload_size_bytes=len(raw_bytes),
                tenant_id=project.tenant_id,
            )
        except QuotaExceededError:
            raise

        document_hash = self.bytes_hash(raw_bytes)
        duplicate = self.find_duplicate_document(
            project_id=project.id,
            document_hash=document_hash,
        )

        storage_key = get_storage_backend().save(
            io.BytesIO(raw_bytes),
            f"{project_id}_{uuid.uuid4().hex[:8]}_{filename}",
        )
        extraction = self.extract_document(filename=filename, raw_bytes=raw_bytes, content_type=file_storage.content_type)
        text_content = extraction.text
        rag_document_id = self.build_rag_document_id(project.id, filename)

        document = ResearcherDocument(
            project_id=project.id,
            filename=filename,
            file_path=storage_key,
            mime_type=file_storage.content_type or "application/octet-stream",
            text_content=text_content,
            file_size=len(raw_bytes),
            status="ready" if text_content else "pending",
            rag_document_id=rag_document_id,
            rag_collection_id=project.collection_id,
            rag_content_hash=self.content_hash(text_content),
            rag_sync_status="not_indexed" if text_content else "unavailable",
            rag_sync_message=None if text_content else "No text could be extracted for AI search.",
            document_hash=document_hash,
        )
        self.apply_extraction_result(document, extraction)
        db.session.add(document)
        db.session.commit()
        self.record_ingestion_state(document, duplicate_of=duplicate)

        quota_service.record_upload(
            user_id=user_id,
            file_size_bytes=len(raw_bytes),
            tenant_id=project.tenant_id,
        )

        if duplicate and duplicate.rag_sync_status == "indexed":
            self._set_rag_status(
                document,
                "skipped_duplicate",
                f"Same content already indexed as document {duplicate.id}.",
            )
            self.record_ingestion_state(document, duplicate_of=duplicate)
            rag_sync = {
                "attempted": False,
                "synced": True,
                "message": document.rag_sync_message,
                "duplicate_of_document_id": duplicate.id,
            }
        else:
            rag_sync = self.sync_document_to_rag(document_id=document.id, user_id=user_id)

        return {
            "document": document,
            "rag_sync": rag_sync,
            "duplicate_of": duplicate,
        }

    def sync_document_to_rag(self, *, document_id: int, user_id: int) -> dict[str, Any]:
        """Retry or perform AI Server RAG sync for a visible document."""

        from app.models.researcher import ResearcherDocument
        from app.services.beep_ai_client import is_configured, sync_document_to_rag

        document = db.session.get(ResearcherDocument, document_id)
        if document is None:
            raise LookupError("Document not found.")
        if document.project_id not in self._visible_project_ids(user_id):
            raise PermissionError("Document is not available to this user.")

        project = document.project
        if not document.text_content:
            self._set_rag_status(document, "unavailable", "Document has no extracted text to index.")
            return {
                "attempted": False,
                "synced": False,
                "message": document.rag_sync_message,
            }
        if not project or not project.collection_id:
            self._set_rag_status(document, "unavailable", "Project is not linked to an AI document library.")
            return {
                "attempted": False,
                "synced": False,
                "message": document.rag_sync_message,
            }
        if not is_configured():
            self._set_rag_status(document, "failed", "Beep.AI.Server is not configured.")
            return {
                "attempted": False,
                "synced": False,
                "message": document.rag_sync_message,
            }

        ok, result = sync_document_to_rag(project, document, user_id=user_id)
        if ok:
            self._set_rag_status(document, "indexed", "Document indexed for AI search.", synced=True)
            self.record_ingestion_state(document)
            return {
                "attempted": True,
                "synced": True,
                "message": document.rag_sync_message,
            }

        self._set_rag_status(document, "failed", str(result))
        self.record_ingestion_state(document, last_error=str(result))
        return {
            "attempted": True,
            "synced": False,
            "message": document.rag_sync_message,
        }

    def extract_text(self, *, filename: str, raw_bytes: bytes) -> str | None:
        """Compatibility wrapper returning only extracted text."""

        return self.extract_document(filename=filename, raw_bytes=raw_bytes).text

    def extract_document(self, *, filename: str, raw_bytes: bytes, content_type: str | None = None):
        """Extract text and metadata through the required parser pipeline."""

        from app.services.document_extraction_service import document_extraction_service

        return document_extraction_service.extract(
            filename=filename,
            raw_bytes=raw_bytes,
            content_type=content_type,
        )

    def get_document_details(self, doc_id: int) -> dict[str, Any]:
        from app.models.researcher import ResearcherDocument
        from app.routes.route_entity_lookup import get_entity_or_404
        from app.services.storage.storage_manager_service import storage_manager_service

        doc = get_entity_or_404(ResearcherDocument, doc_id)
        owner = self.get_document_owner(doc)
        storage_key = self.get_storage_key(doc)
        return {
            "document": doc.to_dict(),
            "owner": {
                "id": owner.id,
                "username": owner.username,
                "email": owner.email,
            } if owner else None,
            "project": doc.project.to_dict() if doc.project else None,
            "storage": {
                "backend": storage_manager_service.backend_name(),
                "reference": self.safe_storage_reference(storage_key),
                "exists": storage_manager_service.object_exists(storage_key),
            },
            "ingestion_state": doc.ingestion_state.to_dict() if getattr(doc, "ingestion_state", None) else None,
        }

    def repair_document(
        self,
        doc_id: int,
        *,
        sync_to_rag: bool = True,
        actor_user_id: int | None = None,
    ) -> dict[str, Any]:
        """Reload a document from storage, rerun extraction, refresh hashes, and optionally sync RAG."""

        from app.models.researcher import ResearcherDocument
        from app.routes.route_entity_lookup import get_entity_or_404
        from app.services.beep_ai_client import is_configured, sync_document_to_rag
        from app.services.storage import get_storage_backend

        doc = get_entity_or_404(ResearcherDocument, doc_id)
        self.record_audit_event(actor_user_id, "admin.documents.repair", doc)
        storage_key = self.get_storage_key(doc)
        raw_bytes = get_storage_backend().load(storage_key)
        document_hash = self.bytes_hash(raw_bytes)
        extraction = self.extract_document(
            filename=doc.filename,
            raw_bytes=raw_bytes,
            content_type=doc.mime_type,
        )
        duplicate = self.find_duplicate_document(
            project_id=doc.project_id,
            document_hash=document_hash,
            exclude_document_id=doc.id,
        )

        doc.file_size = len(raw_bytes)
        doc.document_hash = document_hash
        self.apply_extraction_result(doc, extraction)
        doc.status = "ready" if extraction.has_text else "pending"
        doc.rag_content_hash = self.content_hash(doc.text_content)
        if not extraction.has_text:
            doc.rag_sync_status = "unavailable"
            doc.rag_sync_message = "Document has no extracted text to index."
        elif duplicate and duplicate.rag_sync_status == "indexed":
            doc.rag_sync_status = "skipped_duplicate"
            doc.rag_sync_message = f"Same content already indexed as document {duplicate.id}."
        else:
            doc.rag_sync_status = "not_indexed"
            doc.rag_sync_message = "Document extraction repaired. AI Server sync is pending."
        db.session.commit()
        self.record_ingestion_state(doc, duplicate_of=duplicate)

        rag_sync = {
            "attempted": False,
            "synced": False,
            "message": doc.rag_sync_message,
            "duplicate_of_document_id": getattr(duplicate, "id", None),
        }
        if (
            sync_to_rag
            and extraction.has_text
            and not duplicate
            and doc.project
            and doc.project.collection_id
            and is_configured()
        ):
            owner = self.get_document_owner(doc)
            ok, result = sync_document_to_rag(doc.project, doc, user_id=getattr(owner, "id", None))
            if ok:
                self._set_rag_status(doc, "indexed", "Document indexed for AI search.", synced=True)
                self.record_ingestion_state(doc)
                rag_sync = {"attempted": True, "synced": True, "message": doc.rag_sync_message}
            else:
                self._set_rag_status(doc, "failed", str(result))
                self.record_ingestion_state(doc, last_error=str(result))
                rag_sync = {"attempted": True, "synced": False, "message": doc.rag_sync_message}

        return {
            "document": doc,
            "extraction": extraction,
            "duplicate_of": duplicate,
            "rag_sync": rag_sync,
        }

    def bulk_action(
        self,
        *,
        doc_ids: list[int],
        action: str,
        actor_user_id: int | None = None,
    ) -> dict[str, Any]:
        """Run an admin bulk document action."""

        action = (action or "").strip().lower()
        summary = {
            "action": action,
            "requested": len(doc_ids),
            "succeeded": 0,
            "failed": 0,
            "errors": [],
        }
        for doc_id in doc_ids:
            try:
                if action == "repair":
                    self.repair_document(doc_id, sync_to_rag=True, actor_user_id=actor_user_id)
                elif action == "delete":
                    self.delete_document(doc_id, actor_user_id=actor_user_id)
                elif action == "archive":
                    self.archive_document(doc_id, actor_user_id=actor_user_id)
                elif action == "restore":
                    self.restore_document(doc_id, actor_user_id=actor_user_id)
                else:
                    raise ValueError(f"Unsupported bulk action: {action}")
                summary["succeeded"] += 1
            except Exception as exc:
                db.session.rollback()
                summary["failed"] += 1
                summary["errors"].append({"document_id": doc_id, "error": str(exc)})
        return summary

    def record_admin_job(
        self,
        *,
        user_id: int,
        name: str,
        action: str,
        doc_ids: list[int],
        result: dict[str, Any],
    ):
        """Persist a completed document-manager job using the shared BatchJob model."""

        from app.models.researcher.batch_operations import BatchJob, BatchJobStatus

        failed = int(result.get("failed") or 0)
        succeeded = int(result.get("succeeded") or (1 if failed == 0 else 0))
        total = int(result.get("requested") or len(doc_ids) or succeeded + failed)
        job = BatchJob(
            user_id=user_id,
            name=name,
            description=f"Document manager {action}",
            plugins_config=[],
            data_filters={
                "service": "document_manager",
                "action": action,
                "document_ids": doc_ids,
                "errors": result.get("errors") or [],
            },
            source_data_type="document_manager",
            status=BatchJobStatus.completed if failed == 0 else BatchJobStatus.failed,
            progress=100.0,
            total_records=total,
            processed_records=succeeded + failed,
            successful_records=succeeded,
            failed_records=failed,
            error_message="; ".join(error.get("error", "") for error in result.get("errors") or []) or None,
            started_at=utcnow_naive(),
            completed_at=utcnow_naive(),
        )
        db.session.add(job)
        db.session.commit()
        return job

    def list_admin_jobs(self, *, page: int = 1, per_page: int = 50):
        """List document-manager jobs."""

        from app.models.researcher.batch_operations import BatchJob

        return (
            BatchJob.query
            .filter(BatchJob.source_data_type == "document_manager")
            .order_by(BatchJob.created_at.desc())
            .paginate(page=page, per_page=per_page, error_out=False)
        )

    def delete_document(self, doc_id: int, *, actor_user_id: int | None = None) -> str:
        """Delete document object and DB row, then update quota counters."""

        from app.models.researcher import ResearcherDocument
        from app.routes.route_entity_lookup import get_entity_or_404
        from app.services.beep_ai_client import is_configured, remove_document_from_project_rag
        from app.services.quota_service import quota_service
        from app.services.storage import StorageError, get_storage_backend

        doc = get_entity_or_404(ResearcherDocument, doc_id)
        filename = doc.filename
        file_size = doc.file_size or 0
        owner = self.get_document_owner(doc)
        storage_key = self.get_storage_key(doc)
        self.record_audit_event(actor_user_id, "admin.documents.delete", doc)

        if doc.project and doc.rag_collection_id and is_configured():
            try:
                remove_document_from_project_rag(
                    project=doc.project,
                    document_ids=[doc.rag_document_id or f"researcher_doc_{doc.id}"],
                    user_id=getattr(owner, "id", None),
                )
            except Exception:
                pass

        try:
            get_storage_backend().delete(storage_key)
        except StorageError:
            pass

        db.session.delete(doc)
        db.session.commit()

        if owner:
            try:
                quota_service.record_delete(
                    owner.id,
                    file_size,
                    getattr(doc.project, "tenant_id", None),
                )
            except Exception:
                db.session.rollback()

        return filename

    def archive_document(self, doc_id: int, *, actor_user_id: int | None = None) -> str:
        """Archive a document without deleting its stored object."""

        from app.models.researcher import ResearcherDocument
        from app.routes.route_entity_lookup import get_entity_or_404
        from app.services.beep_ai_client import is_configured, remove_document_from_project_rag

        doc = get_entity_or_404(ResearcherDocument, doc_id)
        self.record_audit_event(actor_user_id, "admin.documents.archive", doc)
        if doc.project and doc.rag_collection_id and is_configured():
            owner = self.get_document_owner(doc)
            try:
                remove_document_from_project_rag(
                    project=doc.project,
                    document_ids=[doc.rag_document_id or f"researcher_doc_{doc.id}"],
                    user_id=getattr(owner, "id", None),
                )
            except Exception:
                pass
        doc.archived_at = utcnow_naive()
        doc.status = "archived"
        doc.rag_sync_status = "archived"
        doc.rag_sync_message = "Document archived by an administrator."
        db.session.commit()
        self.record_ingestion_state(doc)
        return doc.filename

    def restore_document(self, doc_id: int, *, actor_user_id: int | None = None) -> str:
        """Restore an archived document and leave AI Server indexing pending."""

        from app.models.researcher import ResearcherDocument
        from app.routes.route_entity_lookup import get_entity_or_404

        doc = get_entity_or_404(ResearcherDocument, doc_id)
        self.record_audit_event(actor_user_id, "admin.documents.restore", doc)
        doc.archived_at = None
        doc.status = "ready" if doc.text_content else "pending"
        doc.rag_sync_status = "not_indexed" if doc.text_content else "unavailable"
        doc.rag_sync_message = (
            "Document restored. AI Server sync is pending."
            if doc.text_content
            else "Document restored but has no extracted text to index."
        )
        db.session.commit()
        self.record_ingestion_state(doc)
        return doc.filename

    @staticmethod
    def record_audit_event(actor_user_id: int | None, action: str, doc) -> None:
        """Create an audit row for admin document actions when an actor is known."""

        if not actor_user_id:
            return
        from app.models.core import AuditLog

        db.session.add(AuditLog(
            user_id=actor_user_id,
            action=action,
            resource=getattr(doc, "filename", None) or "document",
            resource_id=str(getattr(doc, "id", "")),
            project_id=getattr(doc, "project_id", None),
        ))

    @staticmethod
    def get_document_owner(doc):
        project = getattr(doc, "project", None)
        return getattr(project, "owner", None) if project else None

    @staticmethod
    def _set_rag_status(doc, status: str, message: str, *, synced: bool = False) -> None:
        doc.rag_sync_status = status
        doc.rag_sync_message = message
        doc.rag_collection_id = getattr(doc.project, "collection_id", None)
        doc.rag_document_id = doc.rag_document_id or f"researcher_doc_{doc.id}"
        doc.rag_content_hash = DocumentManagerService.content_hash(doc.text_content)
        if synced:
            doc.rag_synced_at = utcnow_naive()
        db.session.commit()

    @staticmethod
    def find_duplicate_document(*, project_id: int, document_hash: str | None, exclude_document_id: int | None = None):
        if not document_hash:
            return None
        from app.models.researcher import ResearcherDocument

        query = ResearcherDocument.query.filter_by(project_id=project_id, document_hash=document_hash)
        if exclude_document_id:
            query = query.filter(ResearcherDocument.id != exclude_document_id)
        return query.order_by(ResearcherDocument.created_at.asc()).first()

    @staticmethod
    def record_ingestion_state(doc, *, duplicate_of=None, last_error: str | None = None):
        from app.models.researcher import DocumentIngestionState

        document_hash = doc.document_hash or doc.rag_content_hash or f"legacy-document-{doc.id}"
        state = DocumentIngestionState.query.filter_by(document_id=doc.id).first()
        if state is None:
            state = getattr(doc, "ingestion_state", None)
        if state is None:
            state = DocumentIngestionState(document_id=doc.id, project_id=doc.project_id, document_hash=document_hash)
            db.session.add(state)

        state.project_id = doc.project_id
        state.document_hash = document_hash
        state.content_hash = doc.rag_content_hash
        state.rag_document_id = doc.rag_document_id
        state.rag_collection_id = doc.rag_collection_id
        state.extraction_status = doc.extraction_status
        state.rag_sync_status = doc.rag_sync_status
        state.duplicate_of_document_id = getattr(duplicate_of, "id", None)
        state.last_error = last_error or (doc.rag_sync_message if doc.rag_sync_status == "failed" else None)
        state.last_synced_at = doc.rag_synced_at
        if state.duplicate_of_document_id:
            state.ingestion_status = "duplicate"
        elif doc.rag_sync_status == "indexed":
            state.ingestion_status = "synced"
        elif doc.rag_sync_status == "failed":
            state.ingestion_status = "failed"
        elif doc.extraction_status == "ready":
            state.ingestion_status = "extracted"
        elif doc.extraction_status == "failed":
            state.ingestion_status = "failed"
        else:
            state.ingestion_status = "pending"
        db.session.commit()
        return state

    @staticmethod
    def build_rag_document_id(project_id: int, filename: str) -> str:
        seed = f"{project_id}:{filename}:{uuid.uuid4().hex}"
        return f"researcher_doc_{hashlib.sha256(seed.encode('utf-8')).hexdigest()[:24]}"

    @staticmethod
    def content_hash(text_content: str | None) -> str | None:
        if not text_content:
            return None
        return hashlib.sha256(text_content.encode("utf-8")).hexdigest()

    @staticmethod
    def bytes_hash(raw_bytes: bytes | None) -> str | None:
        if not raw_bytes:
            return None
        return hashlib.sha256(raw_bytes).hexdigest()

    @staticmethod
    def apply_extraction_result(doc, extraction) -> None:
        from app.services.document_extraction_service import DocumentExtractionService

        DocumentExtractionService.apply_result_to_document(doc, extraction)

    @staticmethod
    def _visible_project_ids(user_id: int) -> list[int]:
        from app.models.researcher import ProjectMember, ResearchProject

        owned_ids = [
            row[0]
            for row in db.session.query(ResearchProject.id)
            .filter(ResearchProject.owner_id == user_id)
            .all()
        ]
        member_ids = [
            row[0]
            for row in db.session.query(ProjectMember.project_id)
            .filter(ProjectMember.user_id == user_id)
            .all()
        ]
        return sorted(set(owned_ids + member_ids))

    @staticmethod
    def get_storage_key(doc) -> str:
        return getattr(doc, "storage_key", None) or doc.file_path or doc.filename

    @staticmethod
    def safe_storage_reference(storage_key: str | None) -> dict[str, str | None]:
        """Return non-path storage identifiers safe for admin UI display."""

        if not storage_key:
            return {"name": None, "sha256": None}
        key_text = str(storage_key)
        name = key_text.replace("\\", "/").rsplit("/", 1)[-1]
        return {
            "name": name,
            "sha256": hashlib.sha256(key_text.encode("utf-8")).hexdigest(),
        }

    @staticmethod
    def _extract_pdf_text(raw_bytes: bytes) -> str | None:
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw_bytes))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text.strip() or None
        except Exception:
            return None

    @staticmethod
    def _extract_docx_text(raw_bytes: bytes) -> str | None:
        try:
            from docx import Document

            document = Document(io.BytesIO(raw_bytes))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            return text.strip() or None
        except Exception:
            return None

    @staticmethod
    def _extract_xlsx_text(raw_bytes: bytes) -> str | None:
        try:
            from openpyxl import load_workbook

            workbook = load_workbook(io.BytesIO(raw_bytes), read_only=True, data_only=True)
            rows: list[str] = []
            for sheet in workbook.worksheets:
                rows.append(f"Sheet: {sheet.title}")
                for row in sheet.iter_rows(values_only=True):
                    values = [str(value) for value in row if value is not None]
                    if values:
                        rows.append(" | ".join(values))
            return "\n".join(rows).strip() or None
        except Exception:
            return None


document_manager_service = DocumentManagerService()
